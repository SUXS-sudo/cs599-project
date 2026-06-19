from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "CS599_大作业报告.md"
TARGET = ROOT / "docs" / "CS599_大作业报告.docx"


def main() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    headings = collect_headings(lines)
    add_cover_and_toc(doc, headings)
    add_markdown_body(doc, lines)
    add_page_numbers(doc)
    doc.save(TARGET)
    print(TARGET)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color in (
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 14, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 78, 121)),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code.font.size = Pt(9)
        code.paragraph_format.space_after = Pt(3)

    if "TOC Entry" not in styles:
        toc = styles.add_style("TOC Entry", WD_STYLE_TYPE.PARAGRAPH)
        toc.font.name = "Microsoft YaHei"
        toc._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        toc.font.size = Pt(10.5)
        toc.paragraph_format.space_after = Pt(2)


def collect_headings(lines: list[str]) -> list[tuple[int, str, str]]:
    headings = []
    seen: dict[str, int] = {}
    for line in lines:
        match = re.match(r"^(#{1,3})\s+(.+?)\s*$", line)
        if not match:
            continue
        level = len(match.group(1))
        title = clean_inline(match.group(2))
        if title in {"CS599 期末大作业报告", "目录"}:
            continue
        bookmark = make_bookmark_name(title, seen)
        headings.append((level, title, bookmark))
    return headings


def add_cover_and_toc(doc: Document, headings: list[tuple[int, str, str]]) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CS599 期末大作业报告")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("SmartRecipe Multi-Agent System").bold = True

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("面向菜谱问答与个性化饮食推荐的多智能体系统")

    doc.add_paragraph()
    add_metadata_table(doc)
    doc.add_page_break()

    h = doc.add_paragraph(style="Heading 1")
    add_bookmark(h, "目录", "toc")
    h.add_run("目录")
    add_toc_field(doc)
    doc.add_paragraph("提示：若目录页码未显示，在 Word 中右键目录选择“更新域”。下面的目录项可直接 Ctrl+单击跳转。")

    for level, text, bookmark in headings:
        if level > 2:
            continue
        p = doc.add_paragraph(style="TOC Entry")
        p.paragraph_format.left_indent = Inches(0.25 * (level - 1))
        add_internal_hyperlink(p, text, bookmark)
    doc.add_page_break()


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("课程名称", "企业级应用软件设计与开发"),
        ("课程代码", "50120224001 / CS599"),
        ("项目名称", "SmartRecipe Multi-Agent System"),
        ("方向", "方向一：Agentic AI 原生开发"),
        ("学号", "【请填写】"),
        ("姓名", "【请填写】"),
        ("专业", "计算机技术 / 软件工程"),
        ("指导教师", "戚欣"),
        ("提交日期", "2026 年 6 月 22 日"),
        ("GitHub 仓库", "【请填写仓库地址，建议命名为 cs599-project】"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    table._tbl.remove(table.rows[0]._tr)
    format_table(table)


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "请在 Word 中右键更新目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, placeholder, fld_end):
        run = p.add_run()
        run._r.append(el)


def add_markdown_body(doc: Document, lines: list[str]) -> None:
    in_code = False
    code_lang = ""
    pending_table: list[list[str]] = []
    bookmark_seen: dict[str, int] = {}

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            add_table(doc, pending_table)
            pending_table = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lang = line.strip("`").strip()
                if code_lang:
                    p = doc.add_paragraph(style="Code Block")
                    p.add_run(f"[{code_lang}]").bold = True
            else:
                in_code = False
                code_lang = ""
            continue

        if in_code:
            p = doc.add_paragraph(style="Code Block")
            shade_paragraph(p, "F3F6FA")
            p.add_run(line)
            continue

        if not line.strip():
            flush_table()
            continue

        heading = re.match(r"^(#{1,4})\s+(.+?)\s*$", line)
        if heading:
            flush_table()
            level = min(len(heading.group(1)), 3)
            text = clean_inline(heading.group(2))
            if text in {"CS599 期末大作业报告", "目录"}:
                continue
            p = doc.add_paragraph(style=f"Heading {level}")
            bookmark = make_bookmark_name(text, bookmark_seen)
            add_bookmark(p, text, bookmark)
            p.add_run(text)
            continue

        if is_table_line(line):
            row = parse_table_row(line)
            if row and not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row):
                pending_table.append(row)
            continue
        flush_table()

        if line.startswith(">"):
            p = doc.add_paragraph()
            shade_paragraph(p, "F4F6F9")
            p.add_run(clean_inline(line.lstrip("> ").strip()))
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        number = re.match(r"^\d+\.\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bullet.group(1))
        elif number:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, number.group(1))
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            cell.text = clean_inline(row[c]) if c < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)
            if r == 0:
                shade_cell(cell, "E8EEF5")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
    format_table(table)


def format_table(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)


def add_inline_runs(paragraph, text: str) -> None:
    text = text.replace("**", "")
    pattern = re.compile(r"`([^`]+)`")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(clean_inline(text[pos:match.start()]))
        run = paragraph.add_run(match.group(1))
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(clean_inline(text[pos:]))


def is_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def clean_inline(text: str) -> str:
    return re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text).strip()


def make_bookmark_name(text: str, seen: dict[str, int]) -> str:
    base = re.sub(r"[^A-Za-z0-9_]+", "_", text)
    if not base.strip("_"):
        base = "heading"
    base = base[:30].strip("_") or "heading"
    count = seen.get(base, 0)
    seen[base] = count + 1
    return base if count == 0 else f"{base}_{count + 1}"


def add_bookmark(paragraph, text: str, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(abs(hash(name)) % 1000000))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(abs(hash(name)) % 1000000))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(underline)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(rpr)
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def shade_cell(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcpr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    margins = tcpr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margins.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("第 ")
    add_field(p, "PAGE")
    p.add_run(" 页")


def add_field(paragraph, instr: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_text, end):
        run = paragraph.add_run()
        run._r.append(el)


if __name__ == "__main__":
    main()
